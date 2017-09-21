# def datetimeReader():

import numpy as np
from datetime import datetime


from elasticsearch import Elasticsearch
from pwd import esPwdLogin
def esLogin():
    return esPwdLogin()

es = esLogin()



def datetimeReader(timeStr):
# timeStr = '2017-09-06T11:06:28.468035'
    year = int(timeStr.split('-')[0])
    month = int(timeStr.split('-')[1])
    day = int(timeStr.split('-')[2].split('T')[0])
    hour = int(timeStr.split('-')[2].split('T')[1].split(':')[0])
    minute = int(timeStr.split('-')[2].split('T')[1].split(':')[1])
    second = int(timeStr.split('-')[2].split('T')[1].split(':')[2].split('.')[0])
    microsecond = int(timeStr.split('-')[2].split('T')[1].split(':')[2].split('.')[1])
    return datetime(year, month, day, hour, minute, second, microsecond) 


def getExactDistinctNameData(distinctName):
    data = {
        "query" : {
            "constant_score" : {
                "filter" : {
                    "term" : {
                        "distinctName" : distinctName
                        }
                    }
                }
            }
        }
    return data

# def checkExist(index, doc_type ,distinctName):
def checkExist(index, targetCompany, distinctName):
    data = getExactDistinctNameData(distinctName)
    count = es.count(index=index, doc_type=targetCompany, body=data)['count']

    if count >= 1:
        return True
    else:
        return False



def checkDateoutAndDelete(index, targetCompany, distinctName):
    data = getExactDistinctNameData(distinctName)
    outputFilter = ['hits.hits._source.name', 'hits.hits._id', 'hits.hits._type', 'hits.hits._source.distinctName', 'hits.hits._source.url', 'hits.hits._source.createTime', 'hits.hits._score']

    ## deal with first index
    res = es.search(index=index, doc_type=targetCompany, body=data, filter_path=outputFilter)
    ## this is list type, however, there will be only one object(company) in it
    for comp in res['hits']['hits']:
        createTime = datetimeReader(comp['_source']['createTime'])
        nowTime = datetime.utcnow()
        ## basically, the situation will not exist: comp['_source']['distinctName'] != distinctName
        if comp['_source']['distinctName'] == distinctName and (nowTime-createTime).days > 30:
        # if comp['_source']['distinctName'] == distinctName and (nowTime-createTime).seconds > 50:
            print("delete ", distinctName)
            _id = comp['_id']
            es.delete(index=index, doc_type=comp['_type'], id=_id)
            delete = True
        else:
            delete=False

    ## deal with second index(+ _url)
    if delete:
        es.delete_by_query(index=index + '_url', doc_type=targetCompany, body=data)
        return True
    else:
        return False

def directlyDelete(index, targetCompany):
    es.delete_by_query(index=index, doc_type=targetCompany, body={"query":{"match_all":{}}})
    es.delete_by_query(index=index+'_url', doc_type=targetCompany, body={"query":{"match_all":{}}})


def deleteAll():
    if es.indices.exists('companyembedding'):
        es.indices.delete(index='companyembedding')

    if es.indices.exists('companyembedding_url'):
        es.indices.delete(index='companyembedding_url')

    if es.indices.exists('companyembedding_labeled'):
        es.indices.delete(index='companyembedding_labeled')
        
    if es.indices.exists('companyembedding_labeled_url'):
        es.indices.delete(index='companyembedding_labeled_url')

    if es.indices.exists('user_log'):
        es.indices.delete(index='user_log')











import os
from crawlerUtl import getDistinctName

def writeStats(targetCorp, keyWords, outputDir, findingCorpsLi ,labeled):

    data = {
        "query":{
            "match":{
                "info":keyWords
            }
        },
        "highlight":{
            "fields":{
                "info":{}
            }
        }
    }

    outputFilter = ['hits.hits._source.distinctName', 'hits.hits._score', 'hits.hits._source.related' , 'hits.hits.highlight.info']
    if labeled:
        res = es.search(index='companyembedding_labeled', doc_type=targetCorp, body=data, filter_path=outputFilter)
    else:
        res = es.search(index='companyembedding', doc_type=targetCorp, body=data, filter_path=outputFilter)
        
    distinctNames = [(comp['_source']['distinctName'], comp['_score'])  for comp in res['hits']['hits']]



    nowtime = datetime.now()
    filetime = str(nowtime).split()[0].replace("-","") + str(nowtime).split()[1].split(":")[0] + str(nowtime).split()[1].split(":")[1]

    file = open(outputDir + "\\" + targetCorp + filetime + ".txt", 'w', encoding='utf8')
    file.write("Target compaies: " + targetCorp + "\n")
    file.write("Total input: " + str(len(findingCorpsLi)) + " companies" + "\n")
    file.write("Related: " + str((len(distinctNames))) + " companies" + "\n")
    file.write("\n")
    file.write("\n")

    for num, compTuple in enumerate(distinctNames):
        distinctName = compTuple[0]
        file.write(str(num+1) + ". " + distinctName + "\n")
        file.write("score: " + str(compTuple[1]) + "\n")
        data = {
                "query" : {
                    "bool":{
                        "should":[
                                {"match":{"info":keyWords}}, 
                                {"match":{"distinctName":distinctName}}
                        ]}},
                "highlight":{
                    "fields":{"info":{}}
                }}
        outputFilter = ['hits.hits._source.distinctName', 'hits.hits._source.url', 'hits.hits.highlight.info']
        if labeled:
            res = es.search(index='companyembedding_labeled_url', doc_type=targetCorp, body=data, filter_path=outputFilter, size=100)
        else:
            res = es.search(index='companyembedding_url', doc_type=targetCorp, body=data, filter_path=outputFilter, size=100)
            
        for comp in res['hits']['hits']:
            if comp['_source']['distinctName'] == distinctName and comp.get('highlight') != None:
                file.write(comp['_source']['url'] + "\n")
                for info in comp.get('highlight')['info']:
                    file.write(info + "\n")
                file.write("------------------------------------------" + "\n")
        file.write("\n")
        file.write("\n")


from docx import Document
import docx
from docx.shared import RGBColor, Pt
import re
from docx.enum.text import WD_LINE_SPACING

def writeStats_word(targetCorp, keyWords, keywords_emphasize, keywords_filtered, outputDir, findingCorpsLi ,labeled):
    count = es.count(index='companyembedding', doc_type=targetCorp)['count']

    keyWordsData = {"match":{"info":keyWords}}
    keywords_emphasizeData = {"match":{"info":{"query":keywords_emphasize,"boost":2}}}
    keywords_filteredData = {"match":{"info":keywords_filtered}}

    data = {}
    data['size'] = count
    data['query'] = {}
    data['query']['bool'] = {}
    data['query']['bool']['should'] = []
    data['query']['bool']['should'].append(keyWordsData)

    if keywords_emphasize != "":
        data['query']['bool']['should'].append(keywords_emphasizeData)

    if keywords_filtered != "":
        data['query']['bool']['must_not'] = keywords_filteredData

    data['highlight'] = {"fields":{"info":{}}}

    ## must, should: https://www.elastic.co/guide/en/elasticsearch/reference/current/query-dsl-bool-query.html
    ## must not: https://www.elastic.co/guide/en/elasticsearch/guide/1.x/_combining_queries_with_filters.html
    

    outputFilter = ['hits.hits._source.distinctName', 'hits.hits._source.name', 'hits.hits._score', 'hits.hits._source.related' , 'hits.hits.highlight.info']
    if labeled:
        res = es.search(index='companyembedding_labeled', doc_type=targetCorp, body=data, filter_path=outputFilter)
    else:
        res = es.search(index='companyembedding', doc_type=targetCorp, body=data, filter_path=outputFilter)
        
    compTuples = [(comp['_source']['distinctName'], comp['_source']['name'], comp['_score'])  for comp in res['hits']['hits']]

    document = Document()
    document.add_heading(targetCorp, 0)
    document.add_paragraph("Total Input: " + str(len(findingCorpsLi)) + " companies")
    document.add_paragraph("Related: " + str((len(compTuples))) + " companies")
    document.add_paragraph("Keywords: " + ", ".join(keyWords.split()))
    document.add_paragraph("Keywords(Emphasize): " + ", ".join(keywords_emphasize.split()))
    document.add_paragraph("Keywords(Filtered): " + ", ".join(keywords_filtered.split()))
    document.add_paragraph()


    heading = document.add_heading("Ordered List", level=1)
    for num, compTuple in enumerate(compTuples):
        document.add_paragraph(str(num+1) + ". " + compTuple[1] + " (score: " + str(compTuple[2]) + ")")
    document.add_paragraph()







    data = {
            "query" : {
                "bool":{
                    "should":{
                        "match":{"info":keyWords + " " + keywords_emphasize}
                        },
                    # "must":{
                    #     "match":{"distinctName":distinctName}
                    #     }
                    }
                },
            "highlight":{
                "fields":{"info":{}}
            }}



    outputFilter = ['hits.hits._source.distinctName', 'hits.hits._source.url', 'hits.hits.highlight.info']
    if labeled:
        res = es.search(index='companyembedding_labeled_url', doc_type=targetCorp, body=data, filter_path=outputFilter, size=1500)
    else:
        res = es.search(index='companyembedding_url', doc_type=targetCorp, body=data, filter_path=outputFilter, size=1500)
    










    heading = document.add_heading("Detail", level=1)
    for num, compTuple in enumerate(compTuples):
        distinctName = compTuple[0]
        compName = compTuple[1]
        heading = document.add_heading(str(num+1) + ". " + compName + " (score: " + str(compTuple[2]) + ")", level=2)
        heading.line_spacing_rule = WD_LINE_SPACING.DOUBLE

        count = 0
        for comp in res['hits']['hits']:
            if comp['_source']['distinctName'] == distinctName and comp.get('highlight') != None:
                count += 1
                par = document.add_paragraph()
                run = par.add_run("page source "+ str(count) +": ")
                run.bold = True
                hyperlink = addHyperlink(par, comp['_source']['url'])
                
                for num, info in enumerate(comp.get('highlight')['info']):
                    para = document.add_paragraph()
                    run = para.add_run('para' + str(num+1) + '. ')
                    run.bold = True
                    for term in info.split():
                        if re.match(r'^<em>.+</em>$', term):
                            run = para.add_run(term.replace("<em>" , "").replace("</em>" , "") + " ")
                            run.bold = True
                            font = run.font
                            font.color.rgb = RGBColor(255, 0, 0)
                        else:
                            para.add_run(term + " ")
                document.add_paragraph("---------------------------------------------------------")
        document.add_paragraph()
        document.add_paragraph()

    for para in document.paragraphs:
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)

    nowtime = datetime.now()
    filetime = str(nowtime).split()[0].replace("-","") + str(nowtime).split()[1].split(":")[0] + str(nowtime).split()[1].split(":")[1]
    document.save(outputDir + "\\" + targetCorp + filetime + '.docx')



def addHyperlink(paragraph, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    fontcolor = docx.oxml.shared.OxmlElement('w:color')
    fontcolor.set(docx.oxml.shared.qn('w:val'), "0000FF", )
    
    underline = docx.oxml.shared.OxmlElement("w:u")
    underline.set(docx.oxml.shared.qn('w:val'), "single", )

    # Join all the xml elements together add add the required text to the w:r element
    rPr.append(fontcolor)
    rPr.append(underline)
    new_run.append(rPr)
    new_run.text = url
    hyperlink.append(new_run)
    
    # <w:hyperlink>
    #     <w:r>
    #         <w:rPr>
    #             <w:color w:val="0000FF"/>
    #             <w:u w:val="single"/>
    #         </w:rPr>
    #     </w:r>
    # </w:hyperlink>

    paragraph._p.append(hyperlink)
    return hyperlink