import numpy as np
from numpy import dot
from numpy.linalg import norm

# def cosine(keywordVector, docVector, tfidfnorm):
# 	""" related documents j and q are in the concept space by comparing the vectors :
# 		cosine  = ( V1 * V2 ) / ||V1|| x ||V2|| """
# 	return float(dot(keywordVector, docVector) / (norm(keywordVector) * tfidfnorm))

def cosine(vector1, vector2):
	""" related documents j and q are in the concept space by comparing the vectors :
		cosine  = ( V1 * V2 ) / ||V1|| x ||V2|| """
	return float(dot(vector1,vector2) / (norm(vector1) * norm(vector2)))

from crawlerUtl import getDistinctName
import json
import math
import pandas as pd
from collections import Counter


from docx import Document
import docx
from docx.shared import RGBColor, Pt, Inches
import re
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT

from datetime import datetime
from nltk.stem import WordNetLemmatizer
lemmer = WordNetLemmatizer()

def processKeywordLi(keyWords):
    keyWordLi = []
    terms = keyWords.split(', ')
    for term in terms:
        processedTerm = []
        for word in term.split():
            processedTerm.append(lemmer.lemmatize(word))
        keyWordLi.append(" ".join(processedTerm))
    return keyWordLi

def processInfo(info, multiTermKeywords):
    infoTerms = []
    for keyword in multiTermKeywords:
        infoTerms.extend(re.findall(keyword, info))
        info = info.replace(keyword, "")
    infoTerms.extend(info.split())
    return Counter(infoTerms)


def writeStats(targetCorp, keyWords, keywords_emphasize, keywords_filtered, outputDir, findingCorpsLi):
    keyWordLi = processKeywordLi(keyWords)
    keywords_emphasizeLi = processKeywordLi(keywords_emphasize) if keywords_emphasize != "" else []
    keywords_filteredLi = processKeywordLi(keywords_filtered) if keywords_filtered != "" else []
    keywords_positive = keyWordLi + keywords_emphasizeLi
    allKeywords = keyWordLi + keywords_emphasizeLi + keywords_filteredLi

    multiTermKeywords = []
    for keyword in allKeywords:
        if len(keyword.split()) > 1:
            multiTermKeywords.append(keyword)

    companyInfoDir = "C:\\SimCorpFinderData\\companyInfo_v23\\" + targetCorp + "\\companyInfo.json"
    with open(companyInfoDir, 'r', encoding='utf8') as f:
        companyInfoDict = json.loads(f.read())

    # IDF
    docLen = len(companyInfoDict)
    idfDict = {}
    for word in allKeywords:
        f = 0
        for info in companyInfoDict:
            if word in info['info']:       
                f += 1
        idfDict[word] = math.log(docLen / (f + 1))

    keyTfidf = []
    for word in allKeywords:
        if word in keywords_emphasizeLi:
            keyTfidf.append(2 * idfDict[word])
        elif word in keywords_filteredLi:
            keyTfidf.append(-10 * idfDict[word])
        else:
            keyTfidf.append(1 * idfDict[word])

    ##　TF
    companyScoreDict = {}
    companyLenDict = {}
    companyKeywordsDict = {}
        
    ## debug
    # companyInfoLi = []
    for info in companyInfoDict:
        ## debug
        # companyInfoDict = []

        companyInfo = processInfo(info['info'], multiTermKeywords)
        companyScore = 0
        companyInfoLen = sum([times for term, times in companyInfo.items()])
        tfidfLi = []
        tfCount = []
        tfPositiveCount = []
        for word in allKeywords:
            tf = companyInfo.get(word, 0)
            tfidf = tf * idfDict[word]
            tfidfLi.append(tfidf)
            tfCount.append(tf)
            if word in keywords_positive:
                tfPositiveCount.append(tf)
        
        score = cosine(keyTfidf, tfidfLi) * 10 if tfidfLi != [0.0] * len(allKeywords) else 0

        # 如果字出現的次數比較平均，則分數比較高 product(tf/((sum(tf)/len(tf)**len(tf))))
        # if score > 0:
        #     tfPositiveCount = np.array([count for count in tfPositiveCount if count >= 1])
        #     len_keywords_positive = len(keywords_positive)
        #     weight = np.prod(tfPositiveCount) / ((sum(tfPositiveCount)/len_keywords_positive) ** len_keywords_positive)
        #     weight = weight if weight <= 1 else 1
        #     weight = weight ** (1/len(keywords_positive))
        #     score = score * weight

        companyScoreDict[info['name']] = score
        companyLenDict[info['name']] = companyInfoLen
        companyKeywordsDict[info['name']] = dict([(allKeywords[i], tfCount[i])for i in range(len(tfCount))])


        ## debug
    #     companyInfoDict.append(info['name'])
    #     companyInfoDict.append(allKeywords)
    #     companyInfoDict.append(tfCount)
    #     companyInfoDict.append(str(keyTfidf))
    #     companyInfoDict.append(companyInfoLen)
    #     companyInfoDict.append(score)
    #     companyInfoLi.append(companyInfoDict) 
    # from pprint import pprint
    # pprint(sorted(companyInfoLi, key=lambda x: x[5], reverse=True))

    scoreTupleList = sorted(companyScoreDict.items(), key=lambda x: x[1], reverse=True)
    scoreTupleList = [(comp, score) for (comp, score) in scoreTupleList if score != 0.00]


    document = Document()
    document.add_heading(targetCorp, 0)
    par = document.add_paragraph()
    par.add_run("Made By SimCorpFinder (")
    addHyperlink(par, "https://goatwang.github.io/SimCorpFinder/index.html", "20")
    par.add_run("), contact us(linkedin: ")
    addHyperlink(par, "https://www.linkedin.com/in/wanghsuanchung/", "20")
    par.add_run(", Email: jeremy4555@yahoo.com.tw)")
    document.add_paragraph()

    for run in par.runs:    
        run.font.size = Pt(10)


    document.add_heading("Simple Statistics", level=1)
    document.add_paragraph("Total Input: " + str(len(findingCorpsLi)) + " companies")
    document.add_paragraph("Related: " + str((len(scoreTupleList))) + " companies")
    document.add_paragraph("Keywords: " + ", ".join(keyWords.split()))
    document.add_paragraph("Keywords(Emphasize): " + ", ".join(keywords_emphasize.split()))
    document.add_paragraph("Keywords(Filtered): " + ", ".join(keywords_filtered.split()))
    document.add_paragraph()


    heading = document.add_heading("Ordered List", level=1)
    for num, compTuple in enumerate(scoreTupleList):
        document.add_paragraph(str(num+1) + ". " + compTuple[0] + " (score: " + "{:.4f}".format(compTuple[1]) + ")")
    document.add_paragraph()







    heading = document.add_heading("Detail", level=1)
    urlInfoDir = "C:\\SimCorpFinderData\\companyInfo_v23\\" + targetCorp + "\\urlInfo.json"
    with open(urlInfoDir, 'r', encoding='utf8') as f:
        urlInfoDict = json.loads(f.read())

    df = pd.DataFrame(urlInfoDict)
    # for each company
    for num, compTuple in enumerate(scoreTupleList):
        compName = compTuple[0]
        heading = document.add_heading(str(num+1) + ". " + compName + " (score: " + "{:.4f}".format(compTuple[1]) + ")", level=2)
        heading.line_spacing_rule = WD_LINE_SPACING.DOUBLE

        heading = document.add_heading("    (1). Total Terms: " + str(companyLenDict.get(compName)), level=3)
        heading.line_spacing_rule = WD_LINE_SPACING.DOUBLE

        heading = document.add_heading("    (2). All Key Words: ", level=3)
        heading.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        keyworddict = companyKeywordsDict.get(compName)

        for word in allKeywords:
            count = keyworddict.get(word, 0)
            if count != 0:
                if word in keyWordLi:
                    run = heading.add_run("\"" + word + "\"" + ": " + str(count) + ", ")
                    font = run.font
                    font.color.rgb = RGBColor(255, 0, 0)
                elif word in keywords_emphasizeLi:
                    run = heading.add_run("\"" + word + "\"" + ": " + str(count) + ", ")
                    font = run.font
                    font.color.rgb = RGBColor(255, 0, 0)
                else:
                    run = heading.add_run("\"" + word + "\"" + ": " + str(count) + ", ")
                    font = run.font
                    font.color.rgb = RGBColor(0, 255, 0)

        heading = document.add_heading("    (3). Url List: ", level=3)
        heading.line_spacing_rule = WD_LINE_SPACING.DOUBLE

        # for each url
        pageCount = 0
        for row in df[df['name'] == compName].iterrows():
            row = row[1]
            info = row['info']
            # if len(set(allKeywords) & set(info.split())) > 0:
            if len(set(allKeywords) & set(processInfo(info, multiTermKeywords).keys())) > 0:
                pageCount += 1
                heading = document.add_heading("        Url "+ str(pageCount) +": ", level=4)
                heading.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                hyperlink = addHyperlink(heading, row['url'])

                heading = document.add_heading("        Keywords: ", level=4)
                heading.line_spacing_rule = WD_LINE_SPACING.DOUBLE

                keyworddict = companyKeywordsDict.get(compName)

                for word in allKeywords:
                    count = processInfo(info, multiTermKeywords).get(word, 0)

                    if count != 0:
                        par = document.add_paragraph()
                        if word in keyWordLi:
                            run = par.add_run("\t" + word + ": " + str(count))
                            font = run.font
                            font.color.rgb = RGBColor(255, 0, 0)
                        elif word in keywords_emphasizeLi:
                            run = par.add_run("\t" + word + ": " + str(count))
                            font = run.font
                            font.color.rgb = RGBColor(255, 0, 0)
                        else:
                            run = par.add_run("\t" + word + ": " + str(count))
                            font = run.font
                            font.color.rgb = RGBColor(0, 255, 0)

        par = document.add_paragraph()

    for para in document.paragraphs:
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)


    nowtime = datetime.now()
    filetime = str(nowtime).split()[0].replace("-","") + str(nowtime).split()[1].split(":")[0] + str(nowtime).split()[1].split(":")[1]
    document.save(outputDir + "\\" + targetCorp + filetime + '.docx')


def addHyperlink(paragraph, url, font_size="22"):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    fontcolor = docx.oxml.shared.OxmlElement('w:color')
    fontcolor.set(docx.oxml.shared.qn('w:val'), "0000FF", )
    
    underline = docx.oxml.shared.OxmlElement("w:u")
    underline.set(docx.oxml.shared.qn('w:val'), "single", )

    underline = docx.oxml.shared.OxmlElement("w:sz")
    underline.set(docx.oxml.shared.qn('w:val'), font_size, )

    # Join all the xml elements together add add the required text to the w:r element
    rPr.append(fontcolor)
    rPr.append(underline)
    new_run.append(rPr)
    new_run.text = url
    hyperlink.append(new_run)
    
    # <w:hyperlink>
    #     <w:r>
    #         <w:rPr>
    #             <w:color w:val="0000FF"/>
    #             <w:u w:val="single"/>
    #             <w:sz w:val="28"/>
    #         </w:rPr>
    #     </w:r>
    # </w:hyperlink>

    paragraph._p.append(hyperlink)
    return hyperlink